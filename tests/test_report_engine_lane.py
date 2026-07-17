"""Tests for Report Engine lane guard (transversal) and table analyzer."""

from __future__ import annotations

from pathlib import Path

import duckdb
import pytest

from duckclaw.report_engine.lane_guard import (
    assert_docx_uses_report_engine_when_templates_exist,
    actor_has_report_templates,
)
from duckclaw.schema_migrations import run_pending_migrations
from duckclaw.write_command_handlers import dispatch_command


def test_lane_guard_allows_docx_without_templates() -> None:
    assert_docx_uses_report_engine_when_templates_exist(
        blocked_tool="convert_document",
        relative_path="output/notas.md",
        output_format="docx",
        db=None,
    )


def test_lane_guard_allows_any_filename_without_templates(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "empty.duckdb"))
    run_pending_migrations(db)
    # Sin plantillas: pandoc libre (cualquier nombre, cualquier nicho)
    assert_docx_uses_report_engine_when_templates_exist(
        blocked_tool="convert_document",
        relative_path="output/INFORME_MENSUAL_N3.md",
        output_format="docx",
        db=db,
        tenant_id="default",
        actor_email="a@ex.com",
    )


def test_lane_guard_blocks_docx_when_template_registered(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub.duckdb"))
    run_pending_migrations(db)
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_any",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "name": "Plantilla genérica",
            "template_uri": "/vault/cualquier.docx",
            "section_schema": [{"id": "body"}],
            "analyzer_mode": "jinja",
        },
    )
    assert actor_has_report_templates(db, tenant_id="default", actor_email="a@ex.com")
    with pytest.raises(ValueError, match="convert_document bloqueado"):
        assert_docx_uses_report_engine_when_templates_exist(
            blocked_tool="convert_document",
            relative_path="output/cualquier_borrador.md",
            output_format="docx",
            db=db,
            tenant_id="default",
            actor_email="a@ex.com",
        )


def test_lane_guard_escape_ad_hoc_docx(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub2.duckdb"))
    run_pending_migrations(db)
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_x",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "name": "X",
            "template_uri": "/vault/x.docx",
            "section_schema": [{"id": "a"}],
            "analyzer_mode": "jinja",
        },
    )
    assert_docx_uses_report_engine_when_templates_exist(
        blocked_tool="convert_document",
        relative_path="output/carta.md",
        output_format="docx",
        db=db,
        tenant_id="default",
        actor_email="a@ex.com",
        allow_ad_hoc_docx=True,
    )


def test_assert_ready_to_render_blocks_missing_required() -> None:
    from duckclaw.report_engine.render_validate import assert_ready_to_render
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    schema = [
        {"id": "a", "label": "A", "required": True},
        {"id": "b", "label": "B", "required": True},
    ]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="a", content="ok", mode="replace", mark_complete=True)
    with pytest.raises(ValueError, match="required"):
        assert_ready_to_render(state, schema, force=False)
    summary = assert_ready_to_render(state, schema, force=True)
    assert summary["complete_count"] == 1


def test_find_unresolved_placeholders(tmp_path: Path) -> None:
    from docx import Document

    from duckclaw.report_engine.render_validate import find_unresolved_placeholders

    target = tmp_path / "leftover.docx"
    doc = Document()
    doc.add_paragraph("Hola {{ campo.x }} mundo")
    doc.save(str(target))
    left = find_unresolved_placeholders(target)
    assert "campo.x" in left


def test_analyze_table_jinja_with_optional_spaces(tmp_path: Path) -> None:
    from docx import Document

    from duckclaw.report_engine.analyzer import analyze_docx_template

    target = tmp_path / "plantilla_tabla.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "{{campo.1}}"
    table.cell(0, 1).text = "{{ campo.2 }}"
    table.cell(1, 0).text = "Fijo"
    table.cell(1, 1).text = "{{ campo.3 }}"
    doc.save(str(target))

    analysis = analyze_docx_template(target)
    ids = {s["id"] for s in analysis["sections"]}
    assert analysis["analyzer_mode"] == "jinja_tables"
    assert ids == {"campo.1", "campo.2", "campo.3"}
    assert analysis["fields_in_tables"] == 3
    assert analysis["tables"][0]["editable_cells"] == 3
    by_id = {s["id"]: s for s in analysis["sections"]}
    assert by_id["campo.1"]["in_table"] is True
    assert by_id["campo.1"]["table_index"] == 0
