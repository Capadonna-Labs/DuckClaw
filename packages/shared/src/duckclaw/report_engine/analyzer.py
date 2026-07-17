"""Analyze Word templates → section schema (Jinja placeholders, optional headings)."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any

from duckclaw.report_engine.table_analyzer import merge_section_schemas, scan_tables_for_jinja

# Identifiers and dotted paths: {{ body }}, {{ evidencia2.1 }}, {{ejecucion1.2}}
_JINJA_VAR_RE = re.compile(
    r"\{\{\s*([a-zA-Z_][\w]*(?:\.[a-zA-Z0-9_][\w]*)*)\s*\}\}"
)
_WT_TEXT_RE = re.compile(r"<w:t(?:\s[^>]*)?>([^<]*)</w:t>")
_OUTLINE_TITLE_RE = re.compile(
    r"^(?:\d+(?:\.\d+)*[\.\)]\s*)?(?:[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñÁÉÍÓÚÑ0-9\s\-–—]{2,})$"
)


def _read_docx_xml(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as zf:
        raw = zf.read("word/document.xml")
    return raw.decode("utf-8", errors="replace")


def _xml_w_t_plaintext(xml: str) -> str:
    """Concatena textos de w:t para recuperar Jinja partido entre runs de Word."""
    return "".join(_WT_TEXT_RE.findall(xml))


def _sections_from_jinja(xml: str) -> list[dict[str, Any]]:
    plaintext = _xml_w_t_plaintext(xml)
    haystack = f"{plaintext}\n{xml}"
    seen: set[str] = set()
    out: list[dict[str, Any]] = []
    for match in _JINJA_VAR_RE.finditer(haystack):
        var = match.group(1).strip()
        if not var or var in seen:
            continue
        seen.add(var)
        label = var.replace(".", " · ").replace("_", " ").strip()
        out.append(
            {
                "id": var,
                "label": label.title() if label.islower() else label,
                "required": True,
            }
        )
    return out


def _sections_from_headings(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError:
        return []
    doc = Document(str(path))
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for para in doc.paragraphs:
        style = (para.style.name if para.style else "") or ""
        if not style.lower().startswith("heading"):
            continue
        label = (para.text or "").strip()
        if not label:
            continue
        sid = re.sub(r"[^a-zA-Z0-9_]+", "_", label.lower()).strip("_")[:64]
        if not sid or sid in seen:
            continue
        seen.add(sid)
        out.append({"id": sid, "label": label, "required": False})
    return out


def _sections_from_outline_paragraphs(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError:
        return []
    doc = Document(str(path))
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for para in doc.paragraphs:
        label = (para.text or "").strip()
        if not label or len(label) > 120:
            continue
        style = ((para.style.name if para.style else "") or "").lower()
        is_heading = style.startswith("heading") or style.startswith("título")
        looks_title = bool(_OUTLINE_TITLE_RE.match(label)) and (
            label.isupper() or is_heading or len(label.split()) <= 10
        )
        if not looks_title:
            continue
        sid = re.sub(r"[^a-zA-Z0-9_]+", "_", label.lower()).strip("_")[:64]
        if not sid or sid in seen:
            continue
        seen.add(sid)
        out.append({"id": sid, "label": label, "required": False})
    return out


def _analysis_with_tables(
    target: Path,
    *,
    analyzer_mode: str,
    sections: list[dict[str, Any]],
) -> dict[str, Any]:
    table_sections, table_summaries = scan_tables_for_jinja(target)
    merged = merge_section_schemas(sections, table_sections)
    in_table = sum(1 for s in merged if s.get("in_table"))
    mode = analyzer_mode
    if in_table > 0 and analyzer_mode == "jinja":
        mode = "jinja_tables"
    return {
        "analyzer_mode": mode,
        "sections": merged,
        "tables": table_summaries,
        "editable_field_count": len(merged),
        "fields_in_tables": in_table,
    }


def normalize_analyzer_mode_for_storage(mode: str) -> str:
    """Persistido en DuckDB (CHECK jinja|headings|mixed). jinja_tables → jinja."""
    m = (mode or "jinja").strip()
    if m == "jinja_tables":
        return "jinja"
    if m in ("jinja", "headings", "mixed"):
        return m
    return "jinja"


def analyze_docx_template(path: str | Path) -> dict[str, Any]:
    """
    Analiza una plantilla .docx genérica (cualquier nicho).

    Fail-loud: sin placeholders Jinja ni headings/outline detectables → ValueError
    con pista accionable. No inventa secciones ajenas al Word del usuario.
    """
    target = Path(path).expanduser().resolve()
    if not target.is_file():
        raise ValueError(f"Plantilla no encontrada: {target}")
    if target.suffix.lower() != ".docx":
        raise ValueError("Solo se analizan plantillas .docx")

    xml = _read_docx_xml(target)
    jinja_sections = _sections_from_jinja(xml)
    if jinja_sections:
        return _analysis_with_tables(target, analyzer_mode="jinja", sections=jinja_sections)

    heading_sections = _sections_from_headings(target)
    if heading_sections:
        base = _analysis_with_tables(target, analyzer_mode="headings", sections=heading_sections)
        base["warning"] = (
            "Plantilla sin placeholders Jinja detectables en XML. "
            "El render docxtpl no rellenará celdas hasta añadir {{ campo }} en cada hueco editable."
        )
        return base

    outline_sections = _sections_from_outline_paragraphs(target)
    if outline_sections:
        base = _analysis_with_tables(target, analyzer_mode="mixed", sections=outline_sections)
        base["warning"] = (
            "Secciones inferidas por títulos, no por {{ placeholders }}. "
            "Añade Jinja en cada celda editable para preservar tablas al renderizar."
        )
        return base

    raise ValueError(
        "No se detectaron secciones en la plantilla. "
        "Añade huecos Jinja ({{ nombre_seccion }} o {{ grupo.campo }}) "
        "o títulos con estilo Heading 1/2 en Word, y vuelve a registrar."
    )
