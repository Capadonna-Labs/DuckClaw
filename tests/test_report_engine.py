"""Tests for Report Engine core."""

from __future__ import annotations

import json
from pathlib import Path

import duckdb

from duckclaw.report_engine.analyzer import analyze_docx_template
from duckclaw.report_engine.preview import render_preview_html
from duckclaw.report_engine.state import init_state_from_schema, patch_section, summarize_status
from duckclaw.schema_migrations import run_pending_migrations
from duckclaw.write_command_handlers import dispatch_command


def test_analyze_corporate_seed_template() -> None:
    path = (
        Path(__file__).resolve().parents[1]
        / "packages/shared/src/duckclaw/seeds/document_templates/corporate_report.docx"
    )
    if not path.is_file():
        import pytest

        pytest.skip("corporate_report.docx seed missing under packages/shared/src/duckclaw/seeds/document_templates/")
    analysis = analyze_docx_template(path)
    ids = {s["id"] for s in analysis["sections"]}
    assert "title" in ids
    assert "body" in ids


def test_analyze_plain_docx_fails_loud(tmp_path: Path) -> None:
    from docx import Document
    import pytest

    target = tmp_path / "informe_libre.docx"
    doc = Document()
    doc.add_paragraph("Informe mensual de gestión sin placeholders Jinja.")
    doc.save(str(target))
    with pytest.raises(ValueError, match="No se detectaron secciones"):
        analyze_docx_template(target)


def test_analyze_dotted_jinja_placeholders(tmp_path: Path) -> None:
    from docx import Document

    target = tmp_path / "dotted.docx"
    doc = Document()
    doc.add_paragraph("Campo {{ evidencia2.1 }} y {{ body }}")
    doc.save(str(target))
    analysis = analyze_docx_template(target)
    ids = {s["id"] for s in analysis["sections"]}
    assert "evidencia2.1" in ids
    assert "body" in ids
    assert "editable_field_count" in analysis


def test_build_render_context_nests_dotted_ids() -> None:
    from duckclaw.report_engine.state import build_render_context, init_state_from_schema, patch_section

    schema = [{"id": "evidencia2.1", "label": "E2.1"}, {"id": "body", "label": "Body"}]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="evidencia2.1", content="Hecho A", mode="replace")
    state = patch_section(state, section_id="body", content="Cuerpo", mode="replace")
    ctx = build_render_context(state)
    assert ctx["body"] == "Cuerpo"
    # Jinja `{{ evidencia2.1 }}` resuelve el segmento como int; la clave es 1, no '1'.
    assert ctx["evidencia2"][1] == "Hecho A"


def test_init_state_marks_image_kind_and_width() -> None:
    from duckclaw.report_engine.state import init_state_from_schema

    schema = [
        {"id": "intro", "label": "Intro", "kind": "text"},
        {"id": "imagen_1", "label": "Imagen 1", "kind": "image", "width_in": 4.0},
    ]
    state = init_state_from_schema(schema)
    assert state["sections"]["intro"]["kind"] == "text"
    assert "width_in" not in state["sections"]["intro"]
    assert state["sections"]["imagen_1"]["kind"] == "image"
    assert state["sections"]["imagen_1"]["width_in"] == 4.0


def test_image_sections_excluded_from_text_context() -> None:
    from duckclaw.report_engine.state import (
        build_render_context,
        image_render_specs,
        init_state_from_schema,
        patch_section,
    )

    schema = [
        {"id": "intro", "label": "Intro", "kind": "text"},
        {"id": "imagen_1", "label": "Imagen 1", "kind": "image"},
    ]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="intro", content="Hola", mode="replace")
    state = patch_section(state, section_id="imagen_1", content="/vault/a.png", mode="replace")
    ctx = build_render_context(state)
    assert ctx["intro"] == "Hola"
    assert "imagen_1" not in ctx  # se resuelve como InlineImage en el render, no como texto
    specs = image_render_specs(state)
    assert specs == [{"key": "imagen_1", "path": "/vault/a.png", "width_in": 5.5}]


def test_blank_template_schema_has_text_and_image_slots() -> None:
    from duckclaw.report_engine.blank_template import BLANK_SECTION_SCHEMA

    kinds = {s["id"]: s.get("kind") for s in BLANK_SECTION_SCHEMA}
    assert kinds.get("intro") == "text"
    assert kinds.get("imagen_1") == "image"
    assert any(v == "image" for v in kinds.values())
    assert any(v == "text" for v in kinds.values())


def test_fit_inline_image_inches_caps_tall_portrait(tmp_path: Path) -> None:
    from PIL import Image

    from duckclaw.report_engine.render import fit_inline_image_inches

    img = tmp_path / "tall.png"
    Image.new("RGB", (900, 1600), (1, 2, 3)).save(img)
    width, height = fit_inline_image_inches(img, 5.5)
    assert height is not None
    assert height <= 7.0 + 1e-6
    assert width < 5.5
    assert abs(width / height - 900 / 1600) < 0.01


def test_render_tall_image_fits_page_height(tmp_path: Path) -> None:
    import zipfile

    import pytest

    pytest.importorskip("docx")
    pytest.importorskip("docxtpl")
    from PIL import Image

    from duckclaw.report_engine.blank_template import (
        BLANK_SECTION_SCHEMA,
        ensure_blank_template_seed,
    )
    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    output_root = tmp_path / "output"
    output_root.mkdir()
    template_root = tmp_path / "private" / "report_engine"
    template = ensure_blank_template_seed(template_root)

    img_dir = tmp_path / "inbound"
    img_dir.mkdir()
    img_path = img_dir / "tall.png"
    Image.new("RGB", (900, 1600), (40, 50, 60)).save(img_path)

    state = init_state_from_schema(BLANK_SECTION_SCHEMA)
    state = patch_section(state, section_id="titulo", content="Titulo OK", mode="replace")
    state = patch_section(state, section_id="intro", content="Intro", mode="replace")
    state = patch_section(state, section_id="imagen_1", content=str(img_path), mode="replace")

    rendered = render_instance_docx_from_uri(
        template_uri=str(template),
        state_json=json.dumps(state),
        output_root=output_root,
        instance_id="rpt_tall",
        title="Doc alto",
        allowed_roots=[template_root, output_root],
        image_roots=[img_dir, output_root],
    )
    assert rendered.get("images_embedded", 0) >= 1
    with zipfile.ZipFile(rendered["path"]) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    # 7" = 7 * 914400 = 6400800 EMUs
    import re

    extents = re.findall(r'wp:extent[^>]*cy="(\d+)"', xml)
    assert extents, "falta extent de la imagen"
    assert int(extents[0]) <= 6400800 + 1000
    assert "Titulo OK" in xml


def test_render_injects_title_when_titulo_section_empty(tmp_path: Path) -> None:
    import pytest

    pytest.importorskip("docx")
    pytest.importorskip("docxtpl")

    from duckclaw.report_engine.blank_template import (
        BLANK_SECTION_SCHEMA,
        ensure_blank_template_seed,
    )
    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    output_root = tmp_path / "output"
    output_root.mkdir()
    template = ensure_blank_template_seed(tmp_path / "private" / "report_engine")
    state = init_state_from_schema(BLANK_SECTION_SCHEMA)
    state = patch_section(state, section_id="intro", content="Solo intro", mode="replace")

    rendered = render_instance_docx_from_uri(
        template_uri=str(template),
        state_json=json.dumps(state),
        output_root=output_root,
        instance_id="rpt_title_fallback",
        title="Titulo desde instancia",
        allowed_roots=[tmp_path],
        image_roots=[tmp_path],
    )
    from docx import Document

    doc = Document(rendered["path"])
    texts = [p.text for p in doc.paragraphs]
    assert any("Titulo desde instancia" in t for t in texts)


def test_render_blank_document_with_image(tmp_path: Path) -> None:
    import pytest

    pytest.importorskip("docx")
    pytest.importorskip("docxtpl")
    from PIL import Image

    from duckclaw.report_engine.blank_template import (
        BLANK_SECTION_SCHEMA,
        ensure_blank_template_seed,
    )
    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    output_root = tmp_path / "output"
    output_root.mkdir()
    template_root = tmp_path / "private" / "report_engine"
    template = ensure_blank_template_seed(template_root)
    assert template.is_file()
    assert not (output_root / "templates").exists()

    img_dir = tmp_path / "inbound"
    img_dir.mkdir()
    img_path = img_dir / "pic.png"
    Image.new("RGB", (8, 8), (10, 20, 30)).save(img_path)

    state = init_state_from_schema(BLANK_SECTION_SCHEMA)
    state = patch_section(state, section_id="intro", content="Texto de prueba", mode="replace")
    state = patch_section(state, section_id="imagen_1", content=str(img_path), mode="replace")

    rendered = render_instance_docx_from_uri(
        template_uri=str(template),
        state_json=json.dumps(state),
        output_root=output_root,
        instance_id="rpt_blank_test",
        title="Doc en blanco",
        allowed_roots=[template_root, output_root],
        image_roots=[img_dir, output_root],
    )
    assert Path(rendered["path"]).is_file()
    assert Path(rendered["path"]).parent == output_root
    assert rendered["relative_path"] == "Doc_en_blanco_rpt_blank_test.docx"
    assert rendered["byte_size"] > 0
    assert rendered.get("images_embedded", 0) >= 1


def test_markdown_table_collapses_for_docx_cells() -> None:
    from duckclaw.report_engine.docx_content import content_to_docxtpl_value, markdown_tables_to_plain

    md = "| Actividad | Estado |\n|---|---|\n| A | OK |"
    plain = markdown_tables_to_plain(md)
    assert "|" not in plain
    assert "Actividad" in plain and "OK" in plain
    val = content_to_docxtpl_value(md)
    assert val != ""


def test_render_preserves_word_table(tmp_path: Path) -> None:
    from docx import Document

    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    tpl = tmp_path / "tpl.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "{{ field_a }}"
    table.cell(0, 1).text = "Fijo"
    table.cell(1, 0).text = "{{ field_b }}"
    table.cell(1, 1).text = "Pie"
    doc.save(str(tpl))

    schema = [{"id": "field_a", "label": "A"}, {"id": "field_b", "label": "B"}]
    state = init_state_from_schema(schema)
    state = patch_section(
        state,
        section_id="field_a",
        content="Línea 1\nLínea 2",
        mode="replace",
    )
    state = patch_section(state, section_id="field_b", content="Valor B", mode="replace")

    out_root = tmp_path / "out"
    rendered = render_instance_docx_from_uri(
        template_uri=str(tpl),
        state_json=__import__("json").dumps(state),
        output_root=out_root,
        instance_id="rpt_tbl",
        title="Test",
    )
    result = Document(str(rendered["path"]))
    assert len(result.tables) == 1
    cell_text = result.tables[0].cell(0, 0).text
    assert "Línea 1" in cell_text and "Línea 2" in cell_text
    assert "Valor B" in result.tables[0].cell(1, 0).text
    assert "Fijo" in result.tables[0].cell(0, 1).text


def test_render_survives_unfilled_nested_placeholder(tmp_path: Path) -> None:
    """Huecos anidados sin rellenar no deben crashear el render (ChainableUndefined)."""
    import json as _json

    from docx import Document

    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import init_state_from_schema, patch_section

    tpl = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_paragraph("{{ grupo1.item1 }}")
    doc.add_paragraph("{{ grupo2.item1 }}")  # nunca se rellena
    doc.save(str(tpl))

    schema = [{"id": "grupo1.item1", "label": "G1"}, {"id": "grupo2.item1", "label": "G2"}]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="grupo1.item1", content="Relleno G1", mode="replace")

    rendered = render_instance_docx_from_uri(
        template_uri=str(tpl),
        state_json=_json.dumps(state),
        output_root=tmp_path / "out",
        instance_id="rpt_nested",
        title="Test",
    )
    result = Document(str(rendered["path"]))
    body = "\n".join(p.text for p in result.paragraphs)
    assert "Relleno G1" in body
    # El hueco no rellenado quedó vacío, sin '{{' ni excepción.
    assert "{{" not in body


def test_render_numeric_dotted_placeholder_fills_cell(tmp_path: Path) -> None:
    """{{ ejecucion1.1 }} exige clave int 1 en el contexto (bug Jinja str vs int)."""
    import json as _json

    from docx import Document
    from jinja2 import Environment, ChainableUndefined

    from duckclaw.report_engine.render import render_instance_docx_from_uri
    from duckclaw.report_engine.state import build_render_context, init_state_from_schema, patch_section

    state = init_state_from_schema([{"id": "ejecucion1.1", "label": "1.1"}])
    state = patch_section(
        state,
        section_id="ejecucion1.1",
        content="Tabla PostgreSQL preferencias dashboard",
        mode="replace",
    )
    ctx = build_render_context(state)
    assert ctx["ejecucion1"][1] == "Tabla PostgreSQL preferencias dashboard"
    assert (
        Environment(undefined=ChainableUndefined)
        .from_string("{{ ejecucion1.1 }}")
        .render(ctx)
        == "Tabla PostgreSQL preferencias dashboard"
    )

    tpl = tmp_path / "tpl.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{ ejecucion1.1 }}"
    doc.save(str(tpl))

    rendered = render_instance_docx_from_uri(
        template_uri=str(tpl),
        state_json=_json.dumps(state),
        output_root=tmp_path / "out",
        instance_id="rpt_num",
        title="Test",
    )
    result = Document(str(rendered["path"]))
    assert "PostgreSQL" in result.tables[0].cell(0, 0).text


def test_patch_section_append_and_status() -> None:
    schema = [{"id": "obligaciones_1", "label": "Obligaciones 1"}]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="obligaciones_1", content="Reunión X", mode="append")
    state = patch_section(state, section_id="obligaciones_1", content="Entrega Y", mode="append")
    summary = summarize_status(state, schema)
    assert summary["partial_count"] == 1
    state = patch_section(
        state,
        section_id="obligaciones_1",
        content="",
        mode="replace",
        mark_complete=True,
    )
    # replace with empty but mark_complete still partial/empty - test mark complete with content
    state = patch_section(
        state,
        section_id="obligaciones_1",
        content="Cierre",
        mode="replace",
        mark_complete=True,
    )
    summary = summarize_status(state, schema)
    assert summary["complete_count"] == 1


def test_report_engine_migration_and_handlers(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub.duckdb"))
    run_pending_migrations(db)
    schema = [{"id": "intro", "label": "Introducción"}]
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_demo",
            "tenant_id": "default",
            "actor_email": "user@example.com",
            "name": "Demo",
            "template_uri": "/vault/demo.docx",
            "section_schema": schema,
            "analyzer_mode": "jinja",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "create_report_instance",
            "instance_id": "rpt_demo",
            "template_id": "tpl_demo",
            "tenant_id": "default",
            "actor_email": "user@example.com",
            "title": "Informe Junio",
            "period_key": "2026-06",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "patch_report_section",
            "instance_id": "rpt_demo",
            "tenant_id": "default",
            "actor_email": "user@example.com",
            "section_id": "intro",
            "content": "Avance del mes",
            "mode": "replace",
        },
    )
    row = db.execute(
        "SELECT state_json, preview_html FROM main.admin_report_instances WHERE instance_id = 'rpt_demo'"
    ).fetchone()
    assert row is not None
    state = json.loads(str(row[0]))
    assert "Avance" in state["sections"]["intro"]["content"]
    assert "<!DOCTYPE html>" in str(row[1])


def test_soft_delete_report_template_archives_owned_instances(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub_tpl_del.duckdb"))
    run_pending_migrations(db)
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_del_t",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "name": "ToDelete",
            "template_uri": "/vault/t.docx",
            "section_schema": [{"id": "body"}],
            "analyzer_mode": "jinja",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "create_report_instance",
            "instance_id": "rpt_owned",
            "template_id": "tpl_del_t",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "title": "Owned",
            "period_key": "2026-08",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "soft_delete_report_template",
            "template_id": "tpl_del_t",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
        },
    )
    tpl = db.execute(
        "SELECT active FROM main.admin_report_templates WHERE template_id = 'tpl_del_t'"
    ).fetchone()
    inst = db.execute(
        "SELECT active, status FROM main.admin_report_instances WHERE instance_id = 'rpt_owned'"
    ).fetchone()
    assert tpl is not None and tpl[0] is False
    assert inst is not None and inst[0] is False and str(inst[1]) == "archived"


def test_soft_delete_report_instance_allows_new_create(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub_soft_del.duckdb"))
    run_pending_migrations(db)
    schema = [{"id": "body", "label": "Body"}]
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_del",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "name": "Del",
            "template_uri": "/vault/d.docx",
            "section_schema": schema,
            "analyzer_mode": "jinja",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "create_report_instance",
            "instance_id": "rpt_del",
            "template_id": "tpl_del",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "title": "Borrador a eliminar",
        },
    )
    dispatch_command(
        db,
        {
            "command_type": "soft_delete_report_instance",
            "instance_id": "rpt_del",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
        },
    )
    row = db.execute(
        "SELECT active, status FROM main.admin_report_instances WHERE instance_id = 'rpt_del'"
    ).fetchone()
    assert row is not None
    assert row[0] is False
    assert str(row[1]) == "archived"
    dispatch_command(
        db,
        {
            "command_type": "create_report_instance",
            "instance_id": "rpt_del_2",
            "template_id": "tpl_del",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "title": "Nuevo borrador",
        },
    )
    alive = db.execute(
        "SELECT instance_id FROM main.admin_report_instances WHERE instance_id = 'rpt_del_2' AND active = true"
    ).fetchone()
    assert alive is not None


def test_create_allows_multiple_instances_same_template(tmp_path: Path) -> None:
    db = duckdb.connect(str(tmp_path / "hub_multi.duckdb"))
    run_pending_migrations(db)
    schema = [{"id": "body", "label": "Body"}]
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_m",
            "tenant_id": "default",
            "actor_email": "a@ex.com",
            "name": "M",
            "template_uri": "/vault/m.docx",
            "section_schema": schema,
            "analyzer_mode": "jinja",
        },
    )
    for iid, title in (("rpt_1", "Uno"), ("rpt_2", "Dos")):
        dispatch_command(
            db,
            {
                "command_type": "create_report_instance",
                "instance_id": iid,
                "template_id": "tpl_m",
                "tenant_id": "default",
                "actor_email": "a@ex.com",
                "title": title,
            },
        )
    n = db.execute(
        "SELECT count(*) FROM main.admin_report_instances WHERE template_id = 'tpl_m' AND active = true"
    ).fetchone()
    assert n is not None and int(n[0]) == 2


def test_upsert_report_template_blocks_other_owner(tmp_path: Path) -> None:
    import pytest

    db = duckdb.connect(str(tmp_path / "hub_owner.duckdb"))
    run_pending_migrations(db)
    dispatch_command(
        db,
        {
            "command_type": "upsert_report_template",
            "template_id": "tpl_own",
            "tenant_id": "default",
            "actor_email": "owner@ex.com",
            "name": "Mine",
            "template_uri": "/vault/m.docx",
            "section_schema": [{"id": "body"}],
            "analyzer_mode": "jinja",
        },
    )
    with pytest.raises(ValueError, match="otro propietario"):
        dispatch_command(
            db,
            {
                "command_type": "upsert_report_template",
                "template_id": "tpl_own",
                "tenant_id": "default",
                "actor_email": "other@ex.com",
                "name": "Hijack",
                "template_uri": "/vault/h.docx",
                "section_schema": [{"id": "body"}],
                "analyzer_mode": "jinja",
            },
        )


def test_preview_html_renders_sections() -> None:
    state = init_state_from_schema([{"id": "a", "label": "Sección A"}])
    html = render_preview_html(title="T", period_key="2026-06", state=state, section_schema=[{"id": "a"}])
    assert "Sección A" in html


def test_preview_html_renders_image_as_data_uri(tmp_path: Path) -> None:
    from PIL import Image

    img = tmp_path / "shot.png"
    Image.new("RGB", (12, 12), (0, 200, 0)).save(img)
    schema = [
        {"id": "intro", "label": "Intro", "kind": "text"},
        {"id": "imagen_1", "label": "Imagen 1", "kind": "image"},
    ]
    state = init_state_from_schema(schema)
    state = patch_section(state, section_id="intro", content="Hola", mode="replace")
    state = patch_section(
        state, section_id="imagen_1", content=str(img), mode="replace", mark_complete=True
    )
    html = render_preview_html(title="Doc", period_key="", state=state, section_schema=schema)
    assert "data:image/png;base64," in html
    assert str(img) not in html
    assert "shot.png" in html
    assert "<img " in html
