"""Table-aware enrichment for docx template analysis."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

_JINJA_VAR_RE = re.compile(
    r"\{\{\s*([a-zA-Z_][\w]*(?:\.[a-zA-Z0-9_][\w]*)*)\s*\}\}"
)


def _label_for_var(var: str) -> str:
    label = var.replace(".", " · ").replace("_", " ").strip()
    return label.title() if label.islower() else label


def _find_jinja_vars_in_text(text: str) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for match in _JINJA_VAR_RE.finditer(text or ""):
        var = match.group(1).strip()
        if var and var not in seen:
            seen.add(var)
            out.append(var)
    return out


def scan_tables_for_jinja(path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """
    Recorre doc.tables y localiza placeholders Jinja por celda.

    Soporta {{ejecucion1.2}} y {{ ejecucion1.2 }} (espacios opcionales).
    """
    try:
        from docx import Document
    except ImportError:
        return [], []

    doc = Document(str(path))
    sections: list[dict[str, Any]] = []
    seen_ids: set[str] = set()
    table_summaries: list[dict[str, Any]] = []

    for table_index, table in enumerate(doc.tables):
        editable_cells = 0
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell_text = cell.text or ""
                vars_in_cell = _find_jinja_vars_in_text(cell_text)
                if not vars_in_cell:
                    continue
                editable_cells += len(vars_in_cell)
                for var in vars_in_cell:
                    if var in seen_ids:
                        continue
                    seen_ids.add(var)
                    sections.append(
                        {
                            "id": var,
                            "label": _label_for_var(var),
                            "required": True,
                            "in_table": True,
                            "table_index": table_index,
                            "row_index": row_index,
                            "col_index": col_index,
                        }
                    )
        if rows and cols:
            table_summaries.append(
                {
                    "table_index": table_index,
                    "rows": rows,
                    "cols": cols,
                    "editable_cells": editable_cells,
                }
            )

    return sections, table_summaries


def merge_section_schemas(
    xml_sections: list[dict[str, Any]],
    table_sections: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Une secciones del XML plano con metadatos de tabla (prioriza ubicación de tabla)."""
    by_id: dict[str, dict[str, Any]] = {}
    for sec in xml_sections:
        sid = str(sec.get("id") or "").strip()
        if sid:
            by_id[sid] = dict(sec)
    for sec in table_sections:
        sid = str(sec.get("id") or "").strip()
        if not sid:
            continue
        if sid in by_id:
            by_id[sid].update({k: v for k, v in sec.items() if k != "label" or not by_id[sid].get("label")})
        else:
            by_id[sid] = dict(sec)
    # Orden estable: primero orden de xml, luego extras de tabla
    order = [str(s.get("id") or "") for s in xml_sections if s.get("id")]
    for sec in table_sections:
        sid = str(sec.get("id") or "")
        if sid and sid not in order:
            order.append(sid)
    return [by_id[sid] for sid in order if sid in by_id]
