"""Plantilla en blanco (texto + imágenes) para documentos desde cero.

No se versiona un .docx binario: se genera on-demand con python-docx bajo el
vault privado del tenant, no bajo OUTPUT/Drive. El schema marca las secciones
de imagen con kind=image para que el render las inserte como InlineImage.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any

BLANK_TEMPLATE_STEM = "blank_document"
BLANK_IMAGE_SLOTS = 15
# Bump cuando cambia el layout del seed (fuerza regenerar .docx + schema upsert).
BLANK_SCHEMA_VERSION = 2


def _build_blank_schema(slot_count: int = BLANK_IMAGE_SLOTS) -> tuple[list[dict[str, Any]], list[str]]:
    """Schema intercalado imagen_N/texto_N + orden de párrafos del body."""
    n = max(1, min(int(slot_count), BLANK_IMAGE_SLOTS))
    schema: list[dict[str, Any]] = [
        {"id": "titulo", "label": "Título", "kind": "text", "required": False},
        {"id": "intro", "label": "Introducción", "kind": "text", "required": False},
    ]
    body_order: list[str] = ["intro"]
    for i in range(1, n + 1):
        schema.append(
            {
                "id": f"imagen_{i}",
                "label": f"Imagen {i}",
                "kind": "image",
                "required": False,
                "width_in": 5.5,
            }
        )
        schema.append(
            {"id": f"texto_{i}", "label": f"Texto {i}", "kind": "text", "required": False}
        )
        body_order.extend([f"imagen_{i}", f"texto_{i}"])
    schema.append({"id": "cierre", "label": "Cierre", "kind": "text", "required": False})
    body_order.append("cierre")
    return schema, body_order


BLANK_SECTION_SCHEMA, _BODY_ORDER = _build_blank_schema(BLANK_IMAGE_SLOTS)


def _docx_has_placeholder(target: Path, placeholder: str) -> bool:
    """True si el .docx contiene el token Jinja (p. ej. imagen_15)."""
    needle = f"{{{{ {placeholder} }}}}".encode("utf-8")
    alt = f"{{{{{placeholder}}}}}".encode("utf-8")
    try:
        with zipfile.ZipFile(target, "r") as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                data = zf.read(name)
                if needle in data or alt in data:
                    return True
    except Exception:
        return False
    return False


def generate_blank_template_docx(target: Path) -> Path:
    """Genera el .docx en blanco con placeholders {{ id }} para docxtpl."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - entorno sin python-docx
        raise ValueError("python-docx no instalado (uv sync o duckops up)") from exc

    target.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("{{ titulo }}", level=0)
    for section_id in _BODY_ORDER:
        doc.add_paragraph("{{ %s }}" % section_id)
    doc.save(str(target))
    return target


def ensure_blank_template_seed(template_root: Path, *, force: bool = False) -> Path:
    """Devuelve la ruta del .docx en blanco privado, regenerándolo si falta o está obsoleto."""
    target = template_root / "templates" / f"{BLANK_TEMPLATE_STEM}.docx"
    last_slot = f"imagen_{BLANK_IMAGE_SLOTS}"
    if (
        not force
        and target.is_file()
        and target.stat().st_size > 0
        and _docx_has_placeholder(target, last_slot)
    ):
        return target
    return generate_blank_template_docx(target)
